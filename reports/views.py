from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime

def balance_sheet_form(request):
    if request.method == "POST":
        # Collect data from form
        cash = float(request.POST.get("cash", 0))
        ar = float(request.POST.get("ar", 0))
        loans = float(request.POST.get("loans", 0))
        equity_val = float(request.POST.get("equity", 0))

        # Create Word document
        doc = Document()
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run(f"My Company\nBalance Sheet\nAs of {datetime.date.today().strftime('%B %d, %Y')}")
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        def add_section(section_title, items):
            doc.add_paragraph(section_title, style="Heading 2")
            total = 0
            for account, amount in items:
                doc.add_paragraph(f"{account}\t{amount}")
                total += amount
            total_p = doc.add_paragraph()
            total_p.add_run(f"Total {section_title}: ").bold = True
            total_p.add_run(str(total))
            total_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        assets = [("Cash", cash), ("Accounts Receivable", ar)]
        liabilities = [("Loans", loans)]
        equity = [("Owner's Equity", equity_val)]

        add_section("Assets", assets)
        doc.add_paragraph()
        add_section("Liabilities", liabilities)
        doc.add_paragraph()
        add_section("Equity", equity)
        doc.add_paragraph()
        doc.add_paragraph("Prepared by: Admin")

        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = 'attachment; filename="balance_sheet.docx"'
        doc.save(response)
        return response

    return render(request, "balance_sheet_form.html")
